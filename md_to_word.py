import os
import sys

try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_BREAK

    # Add these imports:
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name, guess_lexer
    from pygments.formatters import RawTokenFormatter
    from pygments.token import Token
except ImportError:
    print("Please install python-docx first (e.g. pip install python-docx).")
    sys.exit(1)

def convert_markdown_to_docx(markdown_file_path, output_docx_path):
    """
    Convert a Markdown file to a Word document (DOCX), applying
    basic syntax highlighting styles for code blocks.
    Note: Full color-based syntax highlighting typically requires
    a library like Pygments, but here we'll use a monospaced style.
    """
    doc = Document()

    # Create a 'Code' character style for monospaced code blocks
    styles = doc.styles
    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(10)

    with open(markdown_file_path, "r", encoding="utf-8", errors="ignore") as f:
        lines = f.readlines()

    in_code_block = False
    code_language = ""
    code_buffer = []
    for line in lines:
        line_stripped = line.rstrip("\n")

        # Detect start/end of code blocks: ```lang
        if line_stripped.startswith("```") and not in_code_block:
            in_code_block = True
            code_language = line_stripped[3:].strip()
            code_buffer = []
            continue
        elif line_stripped.startswith("```") and in_code_block:
            # End code block
            _add_code_paragraph(doc, code_buffer, styles["CodeBlock"], code_language)
            in_code_block = False
            code_language = ""
            code_buffer = []
            continue

        if in_code_block:
            code_buffer.append(line_stripped)
        else:
            # Basic check for headings (e.g. "## Some Heading")
            if line_stripped.startswith("##"):
                heading_text = line_stripped.lstrip("#").strip()
                doc.add_heading(heading_text, level=2)
            else:
                doc.add_paragraph(line_stripped)

    doc.save(output_docx_path)
    print(f"Created {output_docx_path} from {markdown_file_path}")

def _highlight_code_paragraph(doc, code_text, style, code_language):
    """
    Highlight code using Pygments tokens, adding each token as a separate run.
    """
    paragraph = doc.add_paragraph(style=style)
    try:
        lexer = get_lexer_by_name(code_language) if code_language else guess_lexer(code_text)
    except:
        lexer = guess_lexer(code_text)

    # Iterate over each (token_type, token_value) from the lexer
    for token_type, token_value in lexer.get_tokens(code_text):
        run = paragraph.add_run(token_value)
        if token_type in Token.Keyword:
            run.font.color.rgb = RGBColor(127, 0, 127)
        elif token_type in Token.String:
            run.font.color.rgb = RGBColor(42, 133, 0)
        elif token_type in Token.Comment:
            run.font.color.rgb = RGBColor(100, 100, 100)

    return paragraph

def _add_code_paragraph(doc, code_lines, style, code_language=""):
    """
    Unify code insertion and highlighting by passing code_lines to _highlight_code_paragraph.
    """
    code_text = "\n".join(code_lines)
    return _highlight_code_paragraph(doc, code_text, style, code_language)

def main():
    # Example usage
    md_file = input("Enter path to the Markdown file: ").strip()
    if not os.path.isfile(md_file):
        print(f"Markdown file not found: {md_file}")
        return

    output_docx = input("Enter the desired output .docx name (e.g. output.docx): ").strip()
    if not output_docx.endswith(".docx"):
        output_docx += ".docx"

    convert_markdown_to_docx(md_file, output_docx)

if __name__ == "__main__":
    main()